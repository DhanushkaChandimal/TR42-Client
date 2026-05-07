"""AI MSA analyst service.

Local-only LLM by default (Ollama). Provider seam in place so a hosted free tier
(e.g. Groq, HF) can be added later without changing call sites.

Helpers (lower in the file):
- text extraction from PDF / DOCX
- HTTP client for Ollama /api/chat with format=json or schema-as-format
- JSON-schema validation for the analyst response

Orchestration (`AiService`):
- can_user_access_msa: tenant scoping so clients never see other clients' MSAs
- analyze_msa: end-to-end run that extracts, calls the model, validates, and
  persists each section as a row in msa_requirement with run_id versioning
- get_analysis / get_pricing: read paths grouped for the frontend
"""
import json
import logging
import os
import uuid
from pathlib import Path

import pdfplumber
import requests
from docx import Document
from flask import current_app
from jsonschema import ValidationError, validate as jsonschema_validate
from sqlalchemy import select

from app.blueprints.enum.enums import UserType
from app.blueprints.repository.msa_repository import MsaRepository
from app.blueprints.repository.msa_requirement_repository import (
    MsaRequirementRepository,
)
from app.blueprints.schema.msa_requirement_schema import (
    msa_requirement_schema,
    msa_requirements_schema,
)
from app.extensions import db
from app.models.client_vendor import ClientVendor
from app.models.msa_requirement import MsaRequirement
from app.models.user import User

logger = logging.getLogger(__name__)


DISCLAIMER_TEXT = (
    "This summary is for informational purposes only and does not constitute legal advice."
)


# JSON shape the model must return. Mirrors the analyst prompt sections plus
# service_rates for pricing extraction. Keys map to msa_requirement.category.
KEY_TERM_TYPES = [
    "parties",
    "effective_date",
    "term_length",
    "renewal",
    "termination",
    "scope",
    "payment_terms",
    "governing_law",
    "jurisdiction",
]
RISK_TYPES = [
    "liability_cap",
    "indemnification",
    "confidentiality",
    "data_protection",
    "ip_ownership",
    "licensing",
    "warranties",
    "disclaimers",
    "insurance",
    "non_compete",
    "non_solicit",
    "exclusivity",
]
RED_FLAG_TYPES = ["one_sided", "ambiguous", "missing"]
ACTION_ITEM_TYPES = ["confirm", "negotiate", "escalate"]


ANALYSIS_JSON_SCHEMA = {
    "type": "object",
    "required": [
        "executive_summary",
        "key_terms",
        "risks",
        "red_flags",
        "action_items",
        "service_rates",
        "disclaimer",
    ],
    "properties": {
        "executive_summary": {
            "type": "object",
            "required": ["summary", "parties", "effective_date", "term_length"],
            "properties": {
                "summary": {"type": "string"},
                "parties": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "required": ["role", "name"],
                        "properties": {
                            "role": {"type": "string"},
                            "name": {"type": "string"},
                            "address": {"type": ["string", "null"]},
                        },
                    },
                },
                "effective_date": {"type": "string"},
                "term_length": {"type": "string"},
                "term_end_or_renewal": {"type": ["string", "null"]},
            },
        },
        "key_terms": {
            "type": "array",
            "items": {
                "type": "object",
                "required": ["rule_type", "description"],
                "properties": {
                    "rule_type": {"type": "string", "enum": KEY_TERM_TYPES},
                    "description": {"type": "string"},
                    "value": {"type": ["string", "null"]},
                    "page_number": {"type": ["integer", "null"]},
                    "extracted_text": {"type": ["string", "null"]},
                    "confidence": {"type": ["number", "null"]},
                },
            },
        },
        "risks": {
            "type": "array",
            "items": {
                "type": "object",
                "required": ["rule_type", "description"],
                "properties": {
                    "rule_type": {"type": "string", "enum": RISK_TYPES},
                    "description": {"type": "string"},
                    "value": {"type": ["string", "null"]},
                    "page_number": {"type": ["integer", "null"]},
                    "extracted_text": {"type": ["string", "null"]},
                    "confidence": {"type": ["number", "null"]},
                },
            },
        },
        "red_flags": {
            "type": "array",
            "items": {
                "type": "object",
                "required": ["rule_type", "description"],
                "properties": {
                    "rule_type": {"type": "string", "enum": RED_FLAG_TYPES},
                    "description": {"type": "string"},
                    "page_number": {"type": ["integer", "null"]},
                    "extracted_text": {"type": ["string", "null"]},
                    "severity": {"type": ["string", "null"]},
                },
            },
        },
        "action_items": {
            "type": "array",
            "items": {
                "type": "object",
                "required": ["rule_type", "description"],
                "properties": {
                    "rule_type": {"type": "string", "enum": ACTION_ITEM_TYPES},
                    "description": {"type": "string"},
                    "priority": {"type": ["string", "null"]},
                },
            },
        },
        "service_rates": {
            "type": "array",
            "items": {
                "type": "object",
                "required": ["service_label", "value", "extracted_text"],
                "properties": {
                    "service_label": {"type": "string"},
                    "value": {"type": "string"},
                    "unit": {"type": ["string", "null"]},
                    "currency": {"type": ["string", "null"]},
                    "page_number": {"type": ["integer", "null"]},
                    "extracted_text": {"type": "string"},
                    "notes": {"type": ["string", "null"]},
                },
            },
        },
        "disclaimer": {"type": "string"},
    },
}


SYSTEM_PROMPT = """You are an AI contract analyst specialized in reviewing Master Service Agreements (MSAs).
You will receive the full text of an MSA. Return a single JSON object that conforms exactly
to the provided schema. Output JSON only. No prose outside the JSON.

GROUNDING RULES (READ CAREFULLY)
- Every finding must come from the contract text provided. Do NOT invent, infer, or copy
  examples from this prompt as if they were real findings.
- If something is not in the document, do NOT make it up. Say "Not specified" or return an
  empty array, depending on the field.
- Any text examples or formats given in this prompt are illustrations of FORMAT only. Never
  treat them as content to include in your output.

CRITICAL RULES
- rule_type must be one of the allowed enum values listed below for each section. Never invent
  new rule_type values. Never use generic words like "phrase", "term", or "clause".
- description is a plain-language explanation of what was found. It is NOT a quotation and
  NOT the rule_type. Two complete sentences max.
- value holds short data only: a number, a percentage, a duration, or a one-line phrase. Long
  verbatim clauses go in extracted_text, not value.
- If a piece of information is not present in the document, say "Not specified" rather than
  guessing or omitting the entry.

executive_summary (object, all fields required):
- summary: 3 to 5 plain-language sentences covering what the agreement covers and the nature
  of the relationship between the parties.
- parties: array of {role, name, address}. Role examples: "Company", "Contractor", "Client",
  "Vendor", "Customer", "Subcontractor". Name is the legal name as written. address is the
  party's address as it appears in the contract, multi-line allowed (use \n for line breaks).
  Set address to null when the contract leaves the field blank or does not list one. If the
  document is a template with placeholders, set name to "Not specified" and address to null.
- effective_date: the date the agreement takes effect, e.g. "2024-01-15" or "January 15, 2024".
  If absent, "Not specified".
- term_length: how long the agreement is valid, e.g. "1 year", "3 years with auto-renewal",
  "month to month". If absent, "Not specified".
- term_end_or_renewal: optional. The end date, renewal terms, or termination notice period.

key_terms: rule_type must be one of:
""" + ", ".join(KEY_TERM_TYPES) + """.
Provide one item per rule_type when the information is present in the document.

risks: rule_type must be one of:
""" + ", ".join(RISK_TYPES) + """.
Include the verbatim clause in extracted_text and page_number when known.

red_flags: rule_type must be one of: one_sided, ambiguous, missing.
Use this section for terms that are atypical, one-sided, ambiguous, or notably absent.

action_items: rule_type must be one of: confirm, negotiate, escalate.
Short list of recommendations for the reviewer.

service_rates: each entry pairs a SERVICE with its RATE as written in the contract.
Do NOT include insurance limits, indemnification caps, or other non-rate dollar amounts.
If no service rates are listed in the contract, return an empty array. NEVER fabricate rates.
extracted_text is REQUIRED and must be a verbatim substring from the document that contains
the service-and-rate pairing. If you cannot quote the rate verbatim from the document, do not
include the entry.
- service_label: the service name as written in the document
- value: the rate amount as a string (numeric portion only, e.g. the dollar amount or count)
- unit: one of per_hour, per_day, per_foot, per_mile, flat, other (or null if not stated)
- currency: extracted from the document if stated, otherwise null
- page_number: the page where the rate appears
- extracted_text: REQUIRED. Verbatim quotation containing the service and its rate.

disclaimer: The exact string "This summary is for informational purposes only and does not constitute legal advice."
"""


class TextExtractionError(Exception):
    pass


class OllamaError(Exception):
    pass


def extract_text(file_path):
    """Extract text from a PDF or DOCX MSA into a list of {page, text} dicts.

    DOCX has no real page boundaries so we synthesize them by paragraph chunks
    so downstream code can still cite a page_number for traceability.
    """
    p = Path(file_path)
    if not p.exists():
        raise TextExtractionError(f"File not found: {file_path}")

    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(p)
    if suffix in (".docx", ".doc"):
        return _extract_docx(p)
    raise TextExtractionError(f"Unsupported file type: {suffix}")


def extract_tables(file_path):
    """Extract tables as structured 2D arrays so the UI can render them as HTML.

    Returns: [{page, rows: [[cell, ...], ...]}, ...]
    Used by the modal to surface every table found in the contract alongside
    the AI's interpretation. The same table content is also injected into the
    text passed to the LLM, but the structured form lets the UI display it
    verbatim as a verification reference.
    """
    p = Path(file_path)
    if not p.exists():
        raise TextExtractionError(f"File not found: {file_path}")

    suffix = p.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf_tables(p)
    if suffix in (".docx", ".doc"):
        return _extract_docx_tables(p)
    return []


def _extract_pdf_tables(path):
    tables = []
    with pdfplumber.open(str(path)) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            try:
                page_tables = page.extract_tables() or []
            except Exception:
                page_tables = []
            for tbl in page_tables:
                rows = [[(c or "").strip().replace("\n", " ") for c in row] for row in tbl]
                rows = [r for r in rows if any(r)]
                if rows:
                    tables.append({"page": idx, "rows": rows})
    return tables


def _extract_docx_tables(path):
    from docx.oxml.ns import qn

    doc = Document(str(path))
    tables = []
    for tbl_el in doc.element.body.iter(qn("w:tbl")):
        rows = []
        for row in tbl_el.iter(qn("w:tr")):
            cells = []
            for cell in row.iter(qn("w:tc")):
                cell_text = " ".join(
                    (t.text or "") for t in cell.iter(qn("w:t"))
                ).strip()
                cells.append(cell_text)
            if any(cells):
                rows.append(cells)
        if rows:
            tables.append({"page": 1, "rows": rows})
    return tables


def _extract_pdf(path):
    """Extract page text + tables from a PDF.

    pdfplumber is used over pypdf because positioned-text tables (a Pricing
    Schedule with rate columns, for example) often don't survive plain text
    extraction. We pull both prose and structured tables and merge them so
    the LLM sees the actual rates rather than the surrounding paragraphs.
    """
    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []
            table_blocks = []
            for tbl in tables:
                rows = []
                for row in tbl:
                    cells = [(c or "").strip().replace("\n", " ") for c in row]
                    if any(cells):
                        rows.append(" | ".join(cells))
                if rows:
                    table_blocks.append("\n".join(rows))
            if table_blocks:
                text = (
                    text
                    + "\n\n--- EXTRACTED TABLES ---\n"
                    + "\n\n".join(table_blocks)
                )
            pages.append({"page": idx, "text": text})
    return pages


def _extract_docx(path):
    """Extract paragraph + table text from a Word document, preserving order.

    python-docx's `doc.paragraphs` only returns body paragraphs and silently
    drops anything inside <w:tbl> blocks. For MSAs the entire pricing schedule
    typically lives in a table, so paragraph-only extraction loses it. We walk
    the body's children in document order, emitting paragraph text and pipe-
    delimited table rows as separate lines, then chunk into synthetic pages.
    """
    from docx.oxml.ns import qn

    doc = Document(str(path))
    body = doc.element.body
    pieces = []
    for child in body:
        if child.tag == qn("w:p"):
            text = "".join((t.text or "") for t in child.iter(qn("w:t")))
            if text.strip():
                pieces.append(text)
        elif child.tag == qn("w:tbl"):
            for row in child.iter(qn("w:tr")):
                cells = []
                for cell in row.iter(qn("w:tc")):
                    cell_text = " ".join(
                        (t.text or "") for t in cell.iter(qn("w:t"))
                    ).strip()
                    cells.append(cell_text)
                if any(cells):
                    pieces.append(" | ".join(cells))

    chunk_size = 30
    pages = []
    for i in range(0, len(pieces), chunk_size):
        pages.append(
            {
                "page": (i // chunk_size) + 1,
                "text": "\n".join(pieces[i : i + chunk_size]),
            }
        )
    return pages or [{"page": 1, "text": ""}]


def pages_to_prompt_text(pages, max_chars=None):
    """Flatten extracted pages into a single tagged string for the prompt."""
    parts = [f"[PAGE {entry['page']}]\n{entry['text']}" for entry in pages]
    text = "\n\n".join(parts)
    if max_chars and len(text) > max_chars:
        text = text[:max_chars]
    return text


class OllamaClient:
    """Thin HTTP client for a local Ollama instance."""

    def __init__(self, base_url=None, model=None, timeout_s=None):
        self.base_url = (
            base_url or os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
        ).rstrip("/")
        self.model = model or os.getenv("OLLAMA_MODEL", "llama3.1:8b")
        self.timeout_s = int(timeout_s or os.getenv("OLLAMA_TIMEOUT_S", "180"))

    def chat_json(self, system_prompt, user_prompt, schema=None):
        """Call /api/chat and return the parsed JSON dict.

        When a schema is provided, Ollama constrains the model's output to match
        it (structured outputs, supported in Ollama 0.5+). Without a schema,
        falls back to plain JSON mode where the model decides the shape.
        """
        url = f"{self.base_url}/api/chat"
        payload = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "stream": False,
            "format": schema if schema else "json",
            "options": {"temperature": 0.1},
        }
        try:
            resp = requests.post(url, json=payload, timeout=self.timeout_s)
        except requests.RequestException as e:
            logger.error(f"Ollama request failed: {e}")
            raise OllamaError(f"Ollama request failed: {e}") from e

        if resp.status_code != 200:
            raise OllamaError(
                f"Ollama returned {resp.status_code}: {resp.text[:500]}"
            )

        body = resp.json()
        content = body.get("message", {}).get("content", "")
        if not content:
            raise OllamaError("Ollama response had empty content")

        try:
            return json.loads(content)
        except json.JSONDecodeError as e:
            raise OllamaError(
                f"Ollama content was not valid JSON: {e}: {content[:500]}"
            ) from e


def validate_analysis(payload):
    """Validate a parsed analysis payload against ANALYSIS_JSON_SCHEMA.

    Returns the payload on success, raises ValueError on schema mismatch.
    """
    try:
        jsonschema_validate(payload, ANALYSIS_JSON_SCHEMA)
    except ValidationError as e:
        raise ValueError(f"AI response failed schema validation: {e.message}") from e
    return payload


class AiService:
    """High-level orchestration for AI MSA analysis."""

    @staticmethod
    def can_user_access_msa(user_id, msa_id):
        """Tenant-scope check. Returns (msa_or_None, has_access).

        - MASTER role: full access.
        - CLIENT users: only MSAs whose vendor is linked to the user's client
          via client_vendor. Prevents client A from seeing client B's contracts.
        - VENDOR / CONTRACTOR / unscoped users: denied for now (no model yet).
        """
        user = User.query.get(user_id)
        if not user:
            return None, False
        msa = MsaRepository.get_by_id(msa_id)
        if not msa:
            return None, False

        role_names = {r.name.upper() for r in user.roles}
        if "MASTER" in role_names:
            return msa, True

        if user.user_type == UserType.CLIENT and user.client_id:
            link = (
                db.session.execute(
                    select(ClientVendor).where(
                        ClientVendor.client_id == user.client_id,
                        ClientVendor.vendor_id == msa.vendor_id,
                    )
                )
                .scalars()
                .first()
            )
            return msa, link is not None

        return msa, False

    @staticmethod
    def _config_value(key, default=None):
        try:
            return current_app.config.get(key, default)
        except RuntimeError:
            return os.getenv(key, default)

    @staticmethod
    def _clip(s, n):
        """Truncate to fit a varchar(n) column. Returns None for None input."""
        if s is None:
            return None
        s = str(s)
        return s if len(s) <= n else s[:n]

    @staticmethod
    def _filter_hallucinated_rates(payload, source_text):
        """Drop service_rates whose quote AND rate value aren't in the source.

        Local LLMs fabricate plausible rates from training data even when the
        contract has none. Two-step check:

        1. extracted_text first 80 chars (normalized) must substring-match the
           document text.
        2. The numeric value (e.g. "150.00", "1,250", "485") must appear
           somewhere in the document. The model can copy real surrounding prose
           into extracted_text but invent a fake price; this catches that.

        Either failure drops the entry. Anything unverifiable is logged and
        dropped before persistence.
        """
        if not source_text:
            return payload
        norm_source = " ".join(source_text.lower().split())
        rates = payload.get("service_rates", []) or []
        grounded = []
        dropped = []
        for entry in rates:
            quote = (entry.get("extracted_text") or "").strip()
            value = (entry.get("value") or "").strip()
            if not quote:
                dropped.append((entry.get("service_label"), "no extracted_text"))
                continue
            snippet = " ".join(quote[:80].lower().split())
            if not (snippet and snippet in norm_source):
                dropped.append((entry.get("service_label"), "quote not in source"))
                continue
            if value:
                # Compare with commas stripped from BOTH sides — source may
                # write "$1,250" while value carries "1250" or vice versa.
                norm_source_no_commas = norm_source.replace(",", "")
                v = value.lower().replace(",", "")
                variants = {v, f"{v}.00", f"${v}"}
                if not any(x in norm_source_no_commas for x in variants if x):
                    dropped.append((entry.get("service_label"), "rate not in source"))
                    continue
            grounded.append(entry)
        if dropped:
            logger.warning(
                f"Dropped {len(dropped)} hallucinated service_rate entries: {dropped}"
            )
        payload["service_rates"] = grounded
        return payload

    @staticmethod
    def _value_or_extract(value, current_extract):
        """Route long strings to extracted_text since value is varchar(100).

        Models routinely return verbatim clauses in the value field. Rather
        than truncating and losing data, promote anything over 100 chars to
        extracted_text (TEXT, unlimited) and leave value None. If
        extracted_text is already populated, keep it as-is.
        """
        if value is None:
            return None, current_extract
        s = str(value)
        if len(s) <= 100:
            return s, current_extract
        return None, current_extract or s

    @staticmethod
    def _build_rows(msa_id, payload, user_id, run_id, model_name, prompt_version):
        """Map a validated analysis payload to a list of MsaRequirement rows."""
        meta_base = {
            "run_id": run_id,
            "is_active": True,
            "model": model_name,
            "prompt_version": prompt_version,
        }
        exec_summary = payload["executive_summary"]
        # The summary text lives in description; structured sub-fields ride in
        # metadata so the UI can render parties/dates/term without re-querying.
        exec_meta = dict(meta_base)
        exec_meta["parties"] = exec_summary.get("parties") or []
        exec_meta["effective_date"] = exec_summary.get("effective_date")
        exec_meta["term_length"] = exec_summary.get("term_length")
        exec_meta["term_end_or_renewal"] = exec_summary.get("term_end_or_renewal")
        rows = [
            MsaRequirement(
                msa_id=msa_id,
                category="executive_summary",
                rule_type="overview",
                description=exec_summary.get("summary", ""),
                extra_metadata=exec_meta,
                created_by=str(user_id),
                updated_by=str(user_id),
            )
        ]
        for kt in payload["key_terms"]:
            v, et = AiService._value_or_extract(kt.get("value"), kt.get("extracted_text"))
            rows.append(
                MsaRequirement(
                    msa_id=msa_id,
                    category="key_term",
                    rule_type=AiService._clip(kt.get("rule_type"), 50),
                    description=kt.get("description"),
                    value=v,
                    page_number=kt.get("page_number"),
                    extracted_text=et,
                    confidence_score=kt.get("confidence"),
                    extra_metadata=dict(meta_base),
                    created_by=str(user_id),
                    updated_by=str(user_id),
                )
            )
        for risk in payload["risks"]:
            v, et = AiService._value_or_extract(risk.get("value"), risk.get("extracted_text"))
            rows.append(
                MsaRequirement(
                    msa_id=msa_id,
                    category="risk",
                    rule_type=AiService._clip(risk.get("rule_type"), 50),
                    description=risk.get("description"),
                    value=v,
                    page_number=risk.get("page_number"),
                    extracted_text=et,
                    confidence_score=risk.get("confidence"),
                    extra_metadata=dict(meta_base),
                    created_by=str(user_id),
                    updated_by=str(user_id),
                )
            )
        for rf in payload["red_flags"]:
            md = dict(meta_base)
            md["severity"] = rf.get("severity")
            rows.append(
                MsaRequirement(
                    msa_id=msa_id,
                    category="red_flag",
                    rule_type=AiService._clip(rf.get("rule_type"), 50),
                    description=rf.get("description"),
                    page_number=rf.get("page_number"),
                    extracted_text=rf.get("extracted_text"),
                    extra_metadata=md,
                    created_by=str(user_id),
                    updated_by=str(user_id),
                )
            )
        for item in payload["action_items"]:
            md = dict(meta_base)
            md["priority"] = item.get("priority")
            rows.append(
                MsaRequirement(
                    msa_id=msa_id,
                    category="action_item",
                    rule_type=AiService._clip(item.get("rule_type"), 50),
                    description=item.get("description"),
                    extra_metadata=md,
                    created_by=str(user_id),
                    updated_by=str(user_id),
                )
            )
        for sr in payload["service_rates"]:
            md = dict(meta_base)
            md["currency"] = sr.get("currency")
            md["notes"] = sr.get("notes")
            v, et = AiService._value_or_extract(sr.get("value"), sr.get("extracted_text"))
            rows.append(
                MsaRequirement(
                    msa_id=msa_id,
                    category="pricing",
                    rule_type="rate",
                    description=sr.get("service_label"),
                    value=v,
                    unit=AiService._clip(sr.get("unit"), 100),
                    page_number=sr.get("page_number"),
                    extracted_text=et,
                    extra_metadata=md,
                    created_by=str(user_id),
                    updated_by=str(user_id),
                )
            )
        return rows

    @staticmethod
    def analyze_msa(msa_id, user_id):
        """Run the full extract -> LLM -> validate -> persist pipeline.

        Returns (response_dict, status_code). Caller is responsible for the
        tenant-scope check before invoking this method.
        """
        msa = MsaRepository.get_by_id(msa_id)
        if not msa:
            return {"message": "MSA not found"}, 404
        if not msa.file_name:
            return {"message": "MSA has no file attached"}, 400

        from app.blueprints.services import storage_service

        try:
            local_path = storage_service.ensure_local(msa.file_name)
        except FileNotFoundError as e:
            return {"message": f"MSA file unavailable: {e}"}, 404
        file_path = str(local_path)

        try:
            pages = extract_text(file_path)
            max_chars = AiService._config_value("AI_MAX_INPUT_CHARS", 200000)
            body = pages_to_prompt_text(pages, max_chars=int(max_chars))
        except TextExtractionError as e:
            return {"message": f"Failed to extract text: {e}"}, 422

        client = OllamaClient(
            base_url=AiService._config_value("OLLAMA_BASE_URL"),
            model=AiService._config_value("OLLAMA_MODEL"),
            timeout_s=AiService._config_value("OLLAMA_TIMEOUT_S", 180),
        )
        try:
            raw = client.chat_json(SYSTEM_PROMPT, body, schema=ANALYSIS_JSON_SCHEMA)
            payload = validate_analysis(raw)
        except OllamaError as e:
            logger.error(f"Ollama call failed for msa {msa_id}: {e}")
            return {"message": f"AI service error: {e}"}, 502
        except ValueError as e:
            logger.error(f"AI response invalid for msa {msa_id}: {e}")
            return {"message": f"AI returned invalid response: {e}"}, 502

        # Defense against the model fabricating service rates from its training
        # data: keep only entries whose extracted_text actually appears in the
        # document we sent.
        payload = AiService._filter_hallucinated_rates(payload, body)

        run_id = str(uuid.uuid4())
        model_name = client.model
        prompt_version = AiService._config_value("AI_PROMPT_VERSION", "v1")

        # Mark prior rows inactive so the reader query (active_only=True) sees
        # only the new run.
        MsaRequirementRepository.deactivate_runs(msa_id)

        rows = AiService._build_rows(
            msa_id=msa_id,
            payload=payload,
            user_id=user_id,
            run_id=run_id,
            model_name=model_name,
            prompt_version=prompt_version,
        )
        MsaRequirementRepository.bulk_create(rows)
        logger.info(f"Analyzed MSA {msa_id} run={run_id} rows={len(rows)}")

        return {
            "msa_id": msa_id,
            "run_id": run_id,
            "row_count": len(rows),
            "disclaimer": DISCLAIMER_TEXT,
            "status": "complete",
        }, 200

    @staticmethod
    def get_analysis(msa_id):
        """Return active-run analysis grouped by section for the UI."""
        rows = MsaRequirementRepository.get_by_msa(msa_id)
        grouped = {
            "executive_summary": None,
            "key_terms": [],
            "risks": [],
            "red_flags": [],
            "action_items": [],
            "pricing": [],
            "disclaimer": DISCLAIMER_TEXT,
            "run_id": None,
        }
        for r in rows:
            data = msa_requirement_schema.dump(r)
            if grouped["run_id"] is None and isinstance(r.extra_metadata, dict):
                grouped["run_id"] = r.extra_metadata.get("run_id")
            if r.category == "executive_summary":
                grouped["executive_summary"] = data
            elif r.category == "key_term":
                grouped["key_terms"].append(data)
            elif r.category == "risk":
                grouped["risks"].append(data)
            elif r.category == "red_flag":
                grouped["red_flags"].append(data)
            elif r.category == "action_item":
                grouped["action_items"].append(data)
            elif r.category == "pricing":
                grouped["pricing"].append(data)
        return grouped, 200

    @staticmethod
    def get_pricing(msa_id):
        rows = MsaRequirementRepository.get_by_msa(msa_id, category="pricing")
        return msa_requirements_schema.dump(rows), 200

    # ------------------------------------------------------------------
    # Team notes (per-MSA, shared across the client's users)
    # ------------------------------------------------------------------

    NOTE_MAX_LEN = 4000

    @staticmethod
    def _serialize_note(row, users_by_id=None):
        author_id = row.created_by
        author_name = None
        if users_by_id and author_id in users_by_id:
            u = users_by_id[author_id]
            full = f"{(u.first_name or '').strip()} {(u.last_name or '').strip()}".strip()
            author_name = full or getattr(u, "username", None) or None
        return {
            "id": row.id,
            "body": row.description or "",
            "created_at": row.created_at.isoformat() if row.created_at else None,
            "author_id": author_id,
            "author_name": author_name,
        }

    @staticmethod
    def add_note(msa_id, body, user_id):
        text = (body or "").strip()
        if not text:
            return {"message": "Note body cannot be empty"}, 400
        if len(text) > AiService.NOTE_MAX_LEN:
            return (
                {"message": f"Note too long (max {AiService.NOTE_MAX_LEN} chars)"},
                400,
            )
        row = MsaRequirementRepository.add_note(msa_id, text, user_id)
        # Resolve the author display name once so the client doesn't
        # have to look it up separately.
        users_by_id = {}
        u = db.session.get(User, str(user_id))
        if u:
            users_by_id[u.id] = u
        return AiService._serialize_note(row, users_by_id), 201

    @staticmethod
    def get_notes(msa_id):
        rows = MsaRequirementRepository.get_notes(msa_id)
        author_ids = {r.created_by for r in rows if r.created_by}
        users_by_id = {}
        if author_ids:
            users = (
                db.session.execute(select(User).where(User.id.in_(author_ids)))
                .scalars()
                .all()
            )
            users_by_id = {u.id: u for u in users}
        return [AiService._serialize_note(r, users_by_id) for r in rows], 200
